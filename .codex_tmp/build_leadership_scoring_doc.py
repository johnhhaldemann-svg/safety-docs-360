from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "leadership-safety-scoring-guide.docx"

NAVY = "102A43"
BLUE = "0E7490"
GREEN = "047857"
AMBER = "B45309"
RED = "B91C1C"
SLATE = "475569"
LIGHT_BLUE = "E0F2FE"
LIGHT_GREEN = "DCFCE7"
LIGHT_AMBER = "FEF3C7"
LIGHT_RED = "FEE2E2"
LIGHT_SLATE = "F1F5F9"
WHITE = "FFFFFF"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color="CBD5E1", size="4"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(cell, top=120, start=140, bottom=120, end=140):
    tc_pr = cell._tc.get_or_add_tcPr()
    margins = tc_pr.find(qn("w:tcMar"))
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        element = margins.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            margins.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")


def set_column_widths(table, widths):
    for row in table.rows:
        for idx, width in enumerate(widths):
            if idx < len(row.cells):
                row.cells[idx].width = Inches(width)


def style_table(table, header_fill=NAVY, header_text=WHITE):
    table.autofit = False
    for row_idx, row in enumerate(table.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell)
            set_cell_border(cell)
            if row_idx == 0:
                set_cell_shading(cell, header_fill)
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.color.rgb = RGBColor.from_string(header_text)
                        run.font.size = Pt(9)
            else:
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.space_after = Pt(0)
                    for run in paragraph.runs:
                        run.font.size = Pt(9)


def add_heading(document, text, level=1):
    paragraph = document.add_heading(text, level=level)
    paragraph.paragraph_format.space_before = Pt(12 if level == 1 else 8)
    paragraph.paragraph_format.space_after = Pt(5)
    for run in paragraph.runs:
        run.font.name = "Aptos Display"
        run.font.color.rgb = RGBColor.from_string(NAVY if level == 1 else BLUE)
    return paragraph


def add_body(document, text):
    paragraph = document.add_paragraph(text)
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.08
    for run in paragraph.runs:
        run.font.name = "Aptos"
        run.font.size = Pt(10.5)
        run.font.color.rgb = RGBColor.from_string("1E293B")
    return paragraph


def add_bullets(document, items):
    for item in items:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(3)
        paragraph.paragraph_format.left_indent = Inches(0.25)
        paragraph.paragraph_format.first_line_indent = Inches(-0.16)
        run = paragraph.add_run("- " + item)
        run.font.name = "Aptos"
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor.from_string("1E293B")


def add_callout(document, title, body, fill=LIGHT_BLUE, title_color=BLUE):
    table = document.add_table(rows=1, cols=1)
    table.autofit = False
    table.columns[0].width = Inches(6.8)
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_border(cell, color="BAE6FD")
    set_cell_margins(cell, top=180, start=220, bottom=180, end=220)
    p1 = cell.paragraphs[0]
    p1.paragraph_format.space_after = Pt(4)
    r1 = p1.add_run(title)
    r1.font.bold = True
    r1.font.size = Pt(10.5)
    r1.font.color.rgb = RGBColor.from_string(title_color)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    r2 = p2.add_run(body)
    r2.font.size = Pt(9.5)
    r2.font.color.rgb = RGBColor.from_string("334155")
    document.add_paragraph().paragraph_format.space_after = Pt(2)


def add_footer(section):
    footer = section.footer.paragraphs[0]
    footer.clear()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("Safety Docs 360 - Leadership Safety Commitment Scoring Guide")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor.from_string(SLATE)


def build():
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    add_footer(section)

    styles = document.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"].font.size = Pt(10.5)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title.paragraph_format.space_after = Pt(4)
    run = title.add_run("Leadership Safety Commitment Scoring Guide")
    run.font.name = "Aptos Display"
    run.font.size = Pt(23)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string(NAVY)

    subtitle = document.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(10)
    run = subtitle.add_run("How Safety Docs 360 calculates evidence-backed leadership scores")
    run.font.name = "Aptos"
    run.font.size = Pt(11.5)
    run.font.color.rgb = RGBColor.from_string(SLATE)

    meta = document.add_table(rows=1, cols=4)
    set_column_widths(meta, [1.5, 1.8, 1.35, 2.15])
    labels = [
        ("Purpose", "Coaching and risk reduction"),
        ("Applies to", "Leadership roles"),
        ("Scale", "0-100, A-F"),
        ("Window", "Default 30 days, configurable 7-120"),
    ]
    for idx, (label, value) in enumerate(labels):
        cell = meta.cell(0, idx)
        set_cell_shading(cell, LIGHT_SLATE)
        set_cell_border(cell)
        set_cell_margins(cell, top=130, start=150, bottom=130, end=150)
        p = cell.paragraphs[0]
        r = p.add_run(label.upper())
        r.font.bold = True
        r.font.size = Pt(7.5)
        r.font.color.rgb = RGBColor.from_string(BLUE)
        p2 = cell.add_paragraph()
        r2 = p2.add_run(value)
        r2.font.bold = True
        r2.font.size = Pt(9)
        r2.font.color.rgb = RGBColor.from_string("0F172A")

    add_heading(document, "Summary", 1)
    add_body(
        document,
        "The leadership score is an automatic safety commitment indicator. It is designed to help company leaders coach risk reduction, not to create a discipline score. The score uses assigned jobsite evidence from permits, JSAs, incidents, corrective actions, AI risk recommendations, and behavior-risk events."
    )
    add_callout(
        document,
        "Plain-language formula",
        "Start at 82 points. Add documented positive signals. Subtract assigned-work risk and follow-through gaps. Clamp the result to 0-100. Convert the final score to an A-F grade and compare it with the previous scoring window for trend.",
        fill=LIGHT_BLUE,
        title_color=BLUE,
    )

    add_heading(document, "Who Is Scored", 1)
    roles = document.add_table(rows=1, cols=3)
    roles.rows[0].cells[0].text = "Role group"
    roles.rows[0].cells[1].text = "Roles"
    roles.rows[0].cells[2].text = "Scoring scope"
    role_rows = [
        ("Company-wide leadership", "Company Admin, Operations Manager, Safety Manager", "All company jobsites"),
        ("Field-scoped leadership", "Project Manager, Field Supervisor, Foreman", "Assigned jobsites only"),
    ]
    for row in role_rows:
        cells = roles.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value
    set_column_widths(roles, [1.65, 2.55, 2.6])
    style_table(roles)

    add_heading(document, "Grade Bands", 1)
    grades = document.add_table(rows=1, cols=3)
    grades.rows[0].cells[0].text = "Score"
    grades.rows[0].cells[1].text = "Grade"
    grades.rows[0].cells[2].text = "Suggested interpretation"
    grade_rows = [
        ("90-100", "A", "Strong evidence of prompt follow-through and low open coaching risk."),
        ("80-89", "B", "Generally strong safety commitment with normal monitoring needs."),
        ("70-79", "C", "Stable but worth coaching on the highest-value open gaps."),
        ("60-69", "D", "Needs focused leadership attention and follow-up planning."),
        ("0-59", "F", "Requires immediate review of assigned-work risk and overdue controls."),
    ]
    for row_idx, row in enumerate(grade_rows, start=1):
        cells = grades.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value
        fill = LIGHT_GREEN if row[1] in {"A", "B"} else LIGHT_BLUE if row[1] == "C" else LIGHT_AMBER if row[1] == "D" else LIGHT_RED
        for cell in cells:
            set_cell_shading(cell, fill)
    set_column_widths(grades, [1.1, 0.8, 4.9])
    style_table(grades)

    document.add_section(WD_SECTION.NEW_PAGE)
    add_heading(document, "Scored Signals", 1)
    add_body(
        document,
        "Signals are sorted by point impact before display. Up to six positive signals, six negative signals, and ten evidence references are retained for a scored leader."
    )
    signals = document.add_table(rows=1, cols=4)
    headers = ["Signal", "Direction", "Point logic", "Evidence source"]
    for idx, header in enumerate(headers):
        signals.rows[0].cells[idx].text = header
    signal_rows = [
        ("No assigned jobsites", "Negative", "-8 when a field-scoped leadership role has no assigned jobsite", "Jobsite assignments"),
        ("Assigned job injury exposure", "Negative", "Severity-based penalty up to -28; recordable and SIF/critical flags add impact; closed status reduces impact", "Incidents"),
        ("Permit closure discipline", "Positive", "+2 per closed permit, capped at +10", "Permits"),
        ("Permit follow-through gaps", "Negative", "Expired/overdue permits, unowned permits, and open stop-work statuses; capped at -24", "Permits"),
        ("JSA process discipline", "Positive", "Closed JSAs count more than active/submitted JSAs; capped at +12", "JSAs"),
        ("JSA quality or closure gaps", "Negative", "Stale drafts and thin descriptions; capped at -24", "JSAs"),
        ("Task control gaps", "Negative", "Weak JSA activity controls or missing permit type detail; capped at -18", "JSA activities"),
        ("Corrective action follow-through", "Positive", "+3 per verified/closed corrective action, capped at +14", "Corrective actions"),
        ("Overdue corrective actions", "Negative", "-5 per overdue open corrective action, capped at -24", "Corrective actions"),
        ("AI risk recommendations available", "Positive", "+2 per active high-confidence recommendation, capped at +8", "AI risk recommendations"),
        ("High-confidence AI recommendations dismissed", "Negative", "-3 per dismissed high-confidence recommendation, capped at -10", "AI risk recommendations"),
        ("Behavior-risk follow-through", "Positive", "+2 per resolved behavior-risk event, capped at +8", "Behavior-risk events"),
        ("Open behavior-risk signals", "Negative", "-3 to -5 per open behavior-risk event depending on driver, capped at -22", "Behavior-risk events"),
        ("No elevated leadership gaps found", "Positive", "+4 when no positive or negative scored signals are found", "Computed fallback"),
    ]
    for row in signal_rows:
        cells = signals.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value
        shade = LIGHT_GREEN if row[1] == "Positive" else LIGHT_AMBER
        set_cell_shading(cells[1], shade)
    set_column_widths(signals, [1.55, 0.9, 2.85, 1.5])
    style_table(signals)

    document.add_section(WD_SECTION.NEW_PAGE)
    add_heading(document, "Data Inputs", 1)
    data = document.add_table(rows=1, cols=3)
    data.rows[0].cells[0].text = "Input"
    data.rows[0].cells[1].text = "Fields used"
    data.rows[0].cells[2].text = "Why it matters"
    data_rows = [
        ("Roles and account status", "user_id, role, team, account_status", "Only active leadership roles are scored."),
        ("Jobsite assignments", "user_id, jobsite_id, role", "Field-scoped leaders are scored on assigned work only."),
        ("Incidents", "severity, status, recordable, injury_type, SIF flag, escalation level", "Captures injury exposure and response closure."),
        ("Permits", "status, due date, owner, stop-work status, SIF flag", "Measures permit discipline and unresolved stop-work risk."),
        ("JSAs and JSA activities", "status, description, mitigation, permit requirement, permit type", "Measures pre-task planning quality and control detail."),
        ("Corrective actions", "status, due date, closed date, severity/priority, SIF potential", "Measures closeout discipline on assigned work."),
        ("AI risk recommendations", "confidence, dismissed status", "Tracks high-confidence risk prompts that need triage."),
        ("Behavior-risk events", "supervisor, driver, status, resolved date", "Tracks supervisor verification, training, permit, and corrective-action signals."),
        ("Previous stored score", "prior score before current window", "Calculates trend in points."),
    ]
    for row in data_rows:
        cells = data.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value
    set_column_widths(data, [1.75, 2.55, 2.5])
    style_table(data)

    add_heading(document, "Visibility Rules", 1)
    add_bullets(
        document,
        [
            "A leader can view their own score.",
            "Company Admins and Operations Managers can view leadership scores across the company.",
            "Safety Managers can view Project Manager, Field Supervisor, and Foreman scores.",
            "Project Managers can view Field Supervisor and Foreman scores when assigned jobsites overlap.",
            "Field Supervisors can view Foreman scores when assigned jobsites overlap.",
            "Scores are persisted to the leadership safety score table when the server has the required admin client.",
        ],
    )

    add_heading(document, "Recommended Use", 1)
    add_bullets(
        document,
        [
            "Use the grade and trend as a coaching opener, not as a stand-alone judgment.",
            "Review the top positive and negative signals before discussing a score.",
            "Open the linked evidence before assigning follow-up work.",
            "Prioritize expired permits, overdue corrective actions, open stop-work items, injury response, and unresolved behavior-risk events.",
            "Recheck the next scoring window after field controls are verified closed.",
        ],
    )

    add_callout(
        document,
        "Implementation note",
        "This document describes the current local scoring model from lib/leadershipSafetyScores.ts, app/api/company/leadership-safety-scores/route.ts, and the related database migration. No platform code or database behavior was changed to create this guide.",
        fill=LIGHT_SLATE,
        title_color=NAVY,
    )

    OUT.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
