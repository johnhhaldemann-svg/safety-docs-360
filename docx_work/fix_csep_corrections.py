from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor


SRC = Path(r"C:\Users\johnh\OneDrive\Desktop\safety_docs_360\docx_work\Em3CiPfK_cover_improved.docx")
OUT = Path(r"C:\Users\johnh\OneDrive\Desktop\safety_docs_360\docx_work\Em3CiPfK_corrected.docx")

PROJECT_NAME = "Candler Ct Structural Steel Project"
FULL_ADDRESS = "N109W15738 Candler Ct, Germantown, WI 53022"
NAVY = "17324D"
TEXT = "252A30"


GLOBAL_REPLACEMENTS = {
    "Test 1": PROJECT_NAME,
    "N109W15738 Candler Ct.": FULL_ADDRESS,
    "N109W15738 Candler Ct": FULL_ADDRESS,
    "Prepared as a presentation draft for review. Replace placeholder logo and project data before final field issue.": (
        "Prepared as a controlled review draft for project verification and signature before field issue."
    ),
    "Current Basis: std federal baseline.": (
        "Current Basis: Federal OSHA construction baseline; no OSHA-approved Wisconsin State Plan overlay for private-sector work."
    ),
    "Clean Citation: OSHA: 29 CFR 1926 Subpart J - Welding, Cutting, Fire Protection, and Prevention.": (
        "Clean Citation: OSHA: 29 CFR 1926 Subpart J - Welding and Cutting; 29 CFR 1926 Subpart F - Fire Protection and Prevention."
    ),
    "Clean Citation: OSHA: 29 CFR 1926.453 - Aerial Lifts and MEWPs.": (
        "Clean Citation: OSHA: 29 CFR 1926.453 - Aerial Lifts (Subpart L); MEWP use per manufacturer instructions and project requirements where applicable."
    ),
    "Clean Citation: OSHA severe weather, heat, cold, lightning, emergency action, and project weather requirements.": (
        "Clean Citation: Project/site severe-weather, heat, cold, lightning, and tornado controls; OSHA 29 CFR 1926.35 - Employee Emergency Action Plans where an EAP is required."
    ),
    "IIPP / Incident Reporting / Corrective Action": "Safety Program / Incident Reporting / Corrective Action",
    "Injury and Illness Prevention Program (IIPP)": "company safety and health program",
    "IIPP": "safety program",
    "Code of Safe Practices": "Safe Work Practices",
    "John Doe / John Doe.": "MJJC project team / MJJC authorized reviewer.",
    "John Doe.": "MJJC project team.",
    "Heaviest Pick: Confirm before issue.": "Heaviest Pick: Per approved lift plan.",
    "Confirm name and phone before field issue.": "Listed on signed field issue contact sheet.",
    "Attach clinic name, address, phone, and route before field issue.": (
        "Listed on signed emergency contact sheet with route before field issue."
    ),
}


PARAGRAPH_REPLACEMENTS = {
    "Current Basis: std federal baseline.": (
        "Current Basis: Federal OSHA construction baseline; no OSHA-approved Wisconsin State Plan overlay for private-sector work."
    ),
    "Clean Citation: OSHA: 29 CFR 1926 Subpart J - Welding, Cutting, Fire Protection, and Prevention.": (
        "Clean Citation: OSHA: 29 CFR 1926 Subpart J - Welding and Cutting; 29 CFR 1926 Subpart F - Fire Protection and Prevention."
    ),
    "Clean Citation: OSHA: 29 CFR 1926.453 - Aerial Lifts and MEWPs.": (
        "Clean Citation: OSHA: 29 CFR 1926.453 - Aerial Lifts (Subpart L); MEWP use per manufacturer instructions and project requirements where applicable."
    ),
    "Clean Citation: OSHA severe weather, heat, cold, lightning, emergency action, and project weather requirements.": (
        "Clean Citation: Project/site severe-weather, heat, cold, lightning, and tornado controls; OSHA 29 CFR 1926.35 - Employee Emergency Action Plans where an EAP is required."
    ),
    "Employees and authorized representatives may examine and receive a copy of the written safety program through supervision, safety, the project office, the company safety office, or the electronic document location used for this CSEP. For California work, access is provided within five business days of a request unless unobstructed electronic access is already available.": (
        "Employees and authorized representatives may request and review applicable written safety-program materials through supervision, the project office, the company safety office, or the electronic document location used for this CSEP. Access is handled according to company policy, contract requirements, and applicable federal OSHA record-access rules."
    ),
    "Records are kept for scheduled and periodic inspections, hazard corrections, training records are maintained for at least one year unless a small-employer exception applies; Cal/OSHA injury and illness logs and related records may have separate retention requirements.": (
        "Records are kept for scheduled and periodic inspections, hazard corrections, training, incidents, permits, and closeout items according to company policy, contract requirements, OSHA 29 CFR 1904 where applicable, and the project document-retention plan."
    ),
    "For California construction work, the company maintains a written Safe Work Practices that applies to the work covered by this CSEP. The Safe Work Practices is posted at a conspicuous location at each jobsite office or provided to each supervisory employee so it is readily available to affected employees.": (
        "The company maintains written safe work practices that apply to the work covered by this CSEP. Applicable practices are made available through supervision, the project office, the company safety office, or the electronic document location used for this CSEP."
    ),
    "Supervisors conduct toolbox, tailgate, or equivalent safety meetings with crews at least every 10 working days for California construction work and more frequently when project conditions require. Meeting topics address current work, incidents, corrective actions, upcoming hazardous tasks, seasonal hazards, and lessons learned; attendance and topics are documented.": (
        "Supervisors conduct toolbox, tailgate, or equivalent safety meetings at the frequency required by company policy, contract, site rules, or changing project conditions. Meeting topics address current work, incidents, corrective actions, upcoming hazardous tasks, seasonal hazards, and lessons learned; attendance and topics are documented."
    ),
    "The safety program includes or cross-references written programs, procedures, forms, and appendices that apply to the work. Common construction appendices include the Safe Work Practices, JHA / PTP forms, inspection forms, incident and near-miss reports, corrective-action logs, training records, emergency contacts, fall protection and rescue procedures, hot-work permits, lift plans, equipment inspection forms, Hazard Communication / SDS information, excavation procedures when triggered, and the task-hazard-control matrix.": (
        "The safety program includes or cross-references written programs, procedures, forms, and appendices that apply to the work. Common construction appendices include safe work practices, JHA / PTP forms, inspection forms, incident and near-miss reports, corrective-action logs, training records, emergency contacts, fall protection and rescue procedures, hot-work permits, lift plans, equipment inspection forms, Hazard Communication / SDS information, excavation procedures when triggered, and the task-hazard-control matrix."
    ),
    "11.14.7 Emergency response location: N109W15738 Candler Ct, Germantown, WI 53022, Germantown, WI 53022": (
        "11.14.7 Emergency response location: N109W15738 Candler Ct, Germantown, WI 53022"
    ),
    "10.2.1 Call 911 for life-threatening injury, fall arrest, fire, collapse, electrocution, severe bleeding, loss of consciousness, suspected fracture, environmental release, or any event requiring outside emergency response; provide the project location as N109W15738 Candler Ct, Germantown, WI 53022, Germantown, WI 53022 R12": (
        "10.2.1 Call 911 for life-threatening injury, fall arrest, fire, collapse, electrocution, severe bleeding, loss of consciousness, suspected fracture, environmental release, or any event requiring outside emergency response; provide the project location as N109W15738 Candler Ct, Germantown, WI 53022. R12"
    ),
    "19.1.1 N109W15738 Candler Ct, Germantown, WI 53022, Germantown, WI 53022": (
        "19.1.1 N109W15738 Candler Ct, Germantown, WI 53022"
    ),
    "training is assigned based on the active scope, permits, equipment, and high-risk programs triggered by this CSEP.": (
        "Training is assigned based on the active scope, permits, equipment, and high-risk programs triggered by this CSEP."
    ),
}


TABLE_CELL_REPLACEMENTS = {
    "Prepared By.": "Prepared By.",
}


APPENDIX_E = {
    "Unload steel.": {
        "Hazards:": "Hazards: struck-by, line of fire, traffic/mobile equipment, pinch/crush points, unstable loads, and overhead load exposure.",
        "Controls:": "Controls: approved delivery route, spotter-supported backing, exclusion zones, load-condition check, stable laydown, and required PPE.",
        "Permits:": "Permits / hold points: delivery and laydown release; lift plan if a crane or hoist is used; elevated-work release only when employees work at elevation.",
        "Training:": "Training: site orientation, material-handling briefing, spotter/traffic-control briefing, and rigging qualification when rigging is assigned.",
        "References:": "References: R1, R5, R6 when hoisting applies, R11, R12, R13",
    },
    "Sort members.": {
        "Hazards:": "Hazards: struck-by, line of fire, pinch/crush points, unstable stockpiles, slips/trips, and equipment interaction.",
        "Controls:": "Controls: sorted laydown by erection sequence, stable blocking, equipment exclusion zones, spotters, housekeeping, and member identification.",
        "Permits:": "Permits / hold points: laydown release; lift plan if a crane or hoist is used; no hot-work permit unless cutting, welding, or grinding is added.",
        "Training:": "Training: site orientation, material-handling briefing, equipment/spotter coordination, and rigging qualification when rigging is assigned.",
        "References:": "References: R1, R5, R6 when hoisting applies, R11, R12, R13",
    },
    "Rigging.": {
        "Hazards:": "Hazards: struck-by, suspended load, dropped load, line of fire, pinch points, rigging failure, and load instability.",
        "Controls:": "Controls: inspected rigging, load weight verification, qualified rigger and signal person, lift plan, barricaded load path, and tag-line control where safe.",
        "Permits:": "Permits / hold points: lift plan or pick plan; critical-lift approval when triggered; elevated-work release if rigging is performed at elevation.",
        "Training:": "Training: qualified rigger, signal person, lift-plan briefing, rigging inspection, and site-specific crane/hoisting coordination.",
        "References:": "References: R1, R6, R11, R12, R13",
    },
    "Crane picks.": {
        "Hazards:": "Hazards: suspended load, swing radius, crane contact or instability, struck-by, load control failure, and wind/weather exposure.",
        "Controls:": "Controls: approved lift plan, operator/rigger/signal-person coordination, load-path barricades, ground-condition verification, weather check, and communication method.",
        "Permits:": "Permits / hold points: lift plan or pick plan; critical-lift approval when triggered; owner / GC / CM hold point where required.",
        "Training:": "Training: crane operator qualification, qualified rigger, signal person, lift-plan briefing, and stop-work/weather trigger briefing.",
        "References:": "References: R1, R6, R12, R13, R17",
    },
    "Column erection.": {
        "Hazards:": "Hazards: fall, overhead work, struck-by, collapse/instability, anchor/fit-up issues, and suspended-load exposure.",
        "Controls:": "Controls: erection sequence verification, anchor and base-condition check, lift plan, temporary stability controls, fall protection, and controlled access below.",
        "Permits:": "Permits / hold points: lift plan; elevated-work release when workers are exposed to fall hazards; anchor/concrete notification before erection starts.",
        "Training:": "Training: steel erection briefing, fall protection and rescue, qualified rigger/signal person where assigned, and competent-person oversight.",
        "References:": "References: R1, R2, R3, R6, R12, R13",
    },
    "Beam setting.": {
        "Hazards:": "Hazards: fall, overhead work, struck-by, collapse/instability, dropped objects, and suspended-load exposure.",
        "Controls:": "Controls: erection sequence verification, lift plan, connector communication, controlled access below, fall protection, and do-not-release stability check.",
        "Permits:": "Permits / hold points: lift plan; elevated-work release when workers are exposed to fall hazards; owner / GC / CM hold point where required.",
        "Training:": "Training: steel erection briefing, fall protection and rescue, connector/crew authorization where applicable, and qualified rigger/signal person where assigned.",
        "References:": "References: R1, R2, R3, R6, R12, R13",
    },
    "Connecting.": {
        "Hazards:": "Hazards: fall, overhead work, struck-by, collapse/instability, dropped objects, and incomplete connection exposure.",
        "Controls:": "Controls: connector authorization, fall protection/rescue readiness, controlled access below, communication before load release, and minimum-connection verification.",
        "Permits:": "Permits / hold points: elevated-work release; lift plan; fall protection/rescue plan; owner / GC / CM hold point where required.",
        "Training:": "Training: connector task briefing, steel erection, fall protection and rescue, qualified signal communication, and stop-work authority.",
        "References:": "References: R1, R2, R3, R6, R12, R13",
    },
    "Bolting.": {
        "Hazards:": "Hazards: fall, overhead work, dropped objects, pinch points, line of fire, tool hazards, and incomplete-connection exposure.",
        "Controls:": "Controls: fall protection, tool inspection, tethering or dropped-object controls where needed, controlled access below, and connection-verification process.",
        "Permits:": "Permits / hold points: elevated-work release when exposed to fall hazards; lift plan only when hoisting supports the task; no hot-work permit unless hot work is added.",
        "Training:": "Training: steel erection briefing, fall protection, hand/power tool safety, dropped-object prevention, and task-specific bolting procedure.",
        "References:": "References: R1, R2, R11, R12, R14",
    },
    "Welding.": {
        "Hazards:": "Hazards: hot work, fire, fumes, burns, arc flash/eye injury, electrical contact, and elevated-work exposure where applicable.",
        "Controls:": "Controls: hot-work permit, fire watch, remove combustibles, spark containment, extinguisher access, ventilation, welding PPE, and protected areas below.",
        "Permits:": "Permits / hold points: hot-work permit; elevated-work release if welding is performed at elevation; owner / GC / CM hot-work approval where required.",
        "Training:": "Training: hot-work/fire-watch training, welder qualification where required, HazCom/SDS review, fire extinguisher use, and fall protection when elevated.",
        "References:": "References: R1, R2 when elevated, R4, R10, R11, R12",
    },
    "Cutting.": {
        "Hazards:": "Hazards: hot work, fire, fumes, burns, sparks/slag, compressed-gas exposure where applicable, and protected-access exposure below.",
        "Controls:": "Controls: hot-work permit, fire watch, remove combustibles, spark containment, cylinder/hose checks where applicable, extinguisher access, and ventilation.",
        "Permits:": "Permits / hold points: hot-work permit; elevated-work release if cutting is performed at elevation; owner / GC / CM hot-work approval where required.",
        "Training:": "Training: hot-work/fire-watch training, cutting equipment authorization, HazCom/SDS review, fire extinguisher use, and fall protection when elevated.",
        "References:": "References: R1, R2 when elevated, R4, R10, R11, R12",
    },
    "Grinding.": {
        "Hazards:": "Hazards: sparks, fire, flying particles, wheel failure, noise, dust/fume exposure, and eye/face injury.",
        "Controls:": "Controls: tool and wheel inspection, guards in place, eye/face and hearing protection, spark containment, combustibles removed, and fire watch when required.",
        "Permits:": "Permits / hold points: hot-work permit when site rules require it for spark-producing work; elevated-work release when performed at elevation.",
        "Training:": "Training: abrasive wheel/tool safety, hot-work/fire-watch training when triggered, PPE use, HazCom/SDS review, and fall protection when elevated.",
        "References:": "References: R1, R2 when elevated, R4 when hot-work rules apply, R10, R11, R12, R14",
    },
    "Decking install.": {
        "Hazards:": "Hazards: fall, collapse/instability, CDZ exposure, falling objects, unsecured bundles, wind/weather exposure, and hot work only when welding/cutting/grinding is performed.",
        "Controls:": "Controls: CDZ or fall-protection controls, fall rescue readiness, decking sequence verification, controlled access below, bundle securement, weather checks, and stability verification before hoisting gear release.",
        "Permits:": "Permits / hold points: elevated-work release; fall protection/rescue plan; lift plan for deck bundles; hot-work permit only if welding, cutting, or grinding is added.",
        "Training:": "Training: decking/CDZ authorization, fall protection and rescue, steel erection briefing, qualified rigging/signal communication when bundles are hoisted, and weather stop-work triggers.",
        "References:": "References: R1, R2, R3, R6 when hoisting applies, R12, R13, R17",
    },
}


def replace_in_runs(container, replacements):
    for paragraph in container.paragraphs:
        for run in paragraph.runs:
            for old, new in replacements.items():
                if old in run.text:
                    run.text = run.text.replace(old, new)
    for table in container.tables:
        for row in table.rows:
            for cell in row.cells:
                replace_in_runs(cell, replacements)


def paragraph_text(paragraph):
    return " ".join(paragraph.text.split())


def set_paragraph_text(paragraph, text):
    if paragraph.runs:
        first = paragraph.runs[0]
        for run in paragraph.runs[1:]:
            run.text = ""
        first.text = text
    else:
        paragraph.add_run(text)


def set_run_font(run, size=9.0, bold=False, color=TEXT, name="Aptos"):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def set_cell_text(cell, text, *, size=9.0, bold=False, color=TEXT):
    paragraph = cell.paragraphs[0]
    for run in paragraph.runs:
        run.text = ""
    run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
    run.text = text
    set_run_font(run, size=size, bold=bold, color=color)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.0
    for extra in cell.paragraphs[1:]:
        for run in extra.runs:
            run.text = ""


def format_cover_table(table):
    for row in table.rows:
        for col_idx, cell in enumerate(row.cells):
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            text = " ".join(cell.text.split())
            set_cell_text(
                cell,
                text,
                size=8.7 if col_idx == 1 else 9.5,
                bold=col_idx == 0,
                color=TEXT if col_idx == 1 else NAVY,
            )


def format_table(table, *, body_size=8.0, header_size=8.0):
    for row_idx, row in enumerate(table.rows):
        for col_idx, cell in enumerate(row.cells):
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            text = " ".join(cell.text.split())
            set_cell_text(
                cell,
                text,
                size=header_size if row_idx == 0 else body_size,
                bold=row_idx == 0 or col_idx == 0,
                color=NAVY if (row_idx == 0 or col_idx == 0) else TEXT,
            )


def apply_exact_paragraph_replacements(doc):
    for paragraph in doc.paragraphs:
        text = paragraph_text(paragraph)
        if text in PARAGRAPH_REPLACEMENTS:
            set_paragraph_text(paragraph, PARAGRAPH_REPLACEMENTS[text])


def apply_table_specifics(doc):
    # Cover table.
    cover = doc.tables[0]
    set_cell_text(cover.rows[0].cells[1], PROJECT_NAME, size=8.7, color=TEXT)
    set_cell_text(cover.rows[1].cells[1], FULL_ADDRESS, size=8.7, color=TEXT)
    format_cover_table(cover)

    # Document control table.
    control = doc.tables[6]
    set_cell_text(control.rows[1].cells[1], PROJECT_NAME, size=7.5)
    set_cell_text(control.rows[5].cells[1], "MJJC project team.", size=7.5)
    set_cell_text(control.rows[6].cells[1], "MJJC authorized reviewer.", size=7.5)
    format_table(control, body_size=7.2, header_size=7.4)

    revision = doc.tables[7]
    set_cell_text(revision.rows[1].cells[3], "MJJC project team / MJJC authorized reviewer.", size=7.0)
    format_table(revision, body_size=6.8, header_size=7.0)

    contacts = doc.tables[15]
    set_cell_text(contacts.rows[2].cells[1], FULL_ADDRESS, size=6.8)
    set_cell_text(contacts.rows[4].cells[1], "Listed on signed field issue contact sheet.", size=6.8)
    set_cell_text(contacts.rows[5].cells[1], "Listed on signed field issue contact sheet.", size=6.8)
    set_cell_text(contacts.rows[6].cells[1], "Listed on signed emergency contact sheet with route before field issue.", size=6.8)
    format_table(contacts, body_size=6.6, header_size=7.0)


def apply_appendix_e(doc):
    active_task = None
    in_appendix_e = False
    for paragraph in doc.paragraphs:
        text = paragraph_text(paragraph)
        if text == "Appendix E. Task-Hazard-Control Matrix":
            in_appendix_e = True
            continue
        if in_appendix_e and text == "Disclaimer":
            in_appendix_e = False
        if in_appendix_e and text:
            is_task = text in APPENDIX_E
            for run in paragraph.runs:
                set_run_font(run, size=8.4 if not is_task else 9.2, bold=False, color=TEXT)
        if text in APPENDIX_E:
            active_task = text
            continue
        if active_task is None:
            continue
        if text.startswith(("Hazards:", "Controls:", "Permits:", "Training:", "References:")):
            prefix = text.split(":", 1)[0] + ":"
            replacement = APPENDIX_E[active_task].get(prefix)
            if replacement:
                set_paragraph_text(paragraph, replacement)
                for run in paragraph.runs:
                    set_run_font(run, size=8.4, bold=False, color=TEXT)
        # Each task block ends after References.
        if text.startswith("References:"):
            active_task = None


def apply_footer(doc):
    for section in doc.sections:
        replace_in_runs(section.footer, {"test |": "Candler Ct |"})


def main():
    doc = Document(SRC)
    replace_in_runs(doc, GLOBAL_REPLACEMENTS)
    apply_exact_paragraph_replacements(doc)
    apply_table_specifics(doc)
    apply_appendix_e(doc)
    apply_footer(doc)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
