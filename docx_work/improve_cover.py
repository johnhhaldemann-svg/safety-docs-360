from io import BytesIO
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


SRC = Path(r"C:\Users\johnh\OneDrive\Desktop\safety_docs_360\docx_work\Em3CiPfK_original.docx")
OUT = Path(r"C:\Users\johnh\OneDrive\Desktop\safety_docs_360\docx_work\Em3CiPfK_cover_improved.docx")

NAVY = "17324D"
BLUE = "2F6F9F"
GOLD = "B07A00"
LIGHT_BLUE = "EAF2F8"
PALE_GOLD = "F8F1DF"
LINE = "B7C8D8"
TEXT = "252A30"
MUTED = "6A737D"


def set_run_font(run, size=None, bold=None, italic=None, color=None, name="Aptos"):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def clear_paragraph(paragraph):
    for run in list(paragraph.runs):
        paragraph._element.remove(run._element)


def set_paragraph_spacing(paragraph, before=0, after=0, line=1.0):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_borders(cell, color=LINE, size="8"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        border = borders.find(qn(tag))
        if border is None:
            border = OxmlElement(tag)
            borders.append(border)
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), size)
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), color)


def set_cell_margins(cell, top=120, start=160, bottom=120, end=160):
    tc_pr = cell._tc.get_or_add_tcPr()
    margins = tc_pr.find(qn("w:tcMar"))
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for side, value in {
        "top": top,
        "start": start,
        "bottom": bottom,
        "end": end,
    }.items():
        node = margins.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, width_twips):
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(width_twips))
    tbl_w.set(qn("w:type"), "dxa")


def style_metadata_table(table):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_width(table, 9360)
    for row_index, row in enumerate(table.rows):
        row.height = Pt(21)
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        for col_index, cell in enumerate(row.cells):
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_borders(cell)
            set_cell_margins(cell, top=65, start=135, bottom=65, end=135)
            set_cell_shading(cell, LIGHT_BLUE if col_index == 0 else "FFFFFF")
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                set_paragraph_spacing(paragraph, before=0, after=0, line=1.05)
                for run in paragraph.runs:
                    set_run_font(
                        run,
                        size=9.5,
                        bold=col_index == 0 or row_index == 0,
                        color=NAVY if col_index == 0 else TEXT,
                    )
            cell.width = Inches(2.15 if col_index == 0 else 4.35)

    for cell in table.rows[0].cells:
        set_cell_shading(cell, PALE_GOLD)


def first_cover_image_blob(doc, paragraph):
    blips = paragraph._element.xpath(".//a:blip")
    if not blips:
        return None
    rid = blips[0].get(qn("r:embed"))
    return doc.part.related_parts[rid].blob if rid in doc.part.related_parts else None


def main():
    doc = Document(SRC)

    # Remove the inherited blank paragraph that pushed the visual center too low.
    body = doc._body._element
    first_child = body[0]
    if first_child.tag == qn("w:p") and not "".join(first_child.xpath(".//w:t/text()")):
        body.remove(first_child)

    image_p, title_p, version_p, desc_p, scope_p, note_p = (
        doc.paragraphs[0],
        doc.paragraphs[1],
        doc.paragraphs[2],
        doc.paragraphs[3],
        doc.paragraphs[4],
        doc.paragraphs[5],
    )
    meta_table = doc.tables[0]

    # Make the existing cover image act as an intentional report banner.
    blob = first_cover_image_blob(doc, image_p)
    clear_paragraph(image_p)
    image_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(image_p, before=0, after=11)
    if blob:
        image_run = image_p.add_run()
        image_run.add_picture(BytesIO(blob), width=Inches(5.05))
    image_p.paragraph_format.keep_with_next = True

    clear_paragraph(title_p)
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(title_p, before=0, after=4)
    title = title_p.add_run("CONTRACTOR SAFETY & ENVIRONMENTAL PLAN")
    set_run_font(title, size=22, bold=True, color=NAVY)
    title_p.paragraph_format.keep_with_next = True

    csep = title_p.add_run(" (CSEP)")
    set_run_font(csep, size=18, bold=True, color=NAVY)

    clear_paragraph(version_p)
    version_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(version_p, before=0, after=5)
    vrun = version_p.add_run("Version C - Reviewer / CODEX Evidence Format")
    set_run_font(vrun, size=13.5, bold=True, italic=True, color=GOLD)
    version_p.paragraph_format.keep_with_next = True

    clear_paragraph(desc_p)
    desc_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(desc_p, before=0, after=2)
    drun = desc_p.add_run("Policy mapping, evidence language, and selective matrices for qualification review.")
    set_run_font(drun, size=10.5, italic=True, color=TEXT)
    desc_p.paragraph_format.keep_with_next = True

    clear_paragraph(scope_p)
    scope_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(scope_p, before=0, after=12)
    srun = scope_p.add_run("Structural Steel / Metals - Steel erection / decking")
    set_run_font(srun, size=10.5, bold=True, color=BLUE)
    scope_p.paragraph_format.keep_with_next = True

    style_metadata_table(meta_table)

    clear_paragraph(note_p)
    note_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(note_p, before=8, after=0)
    nrun = note_p.add_run(
        "Prepared as a presentation draft for review. Replace placeholder logo and project data before final field issue."
    )
    set_run_font(nrun, size=9.5, italic=True, color=MUTED)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
